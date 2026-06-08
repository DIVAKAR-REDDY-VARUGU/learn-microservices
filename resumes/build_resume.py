#!/usr/bin/env python3
"""Reusable resume generator: read a role JSON and emit matching .docx and .pdf.

Usage:
    python build_resume.py <path-to-role.json>

Writes <same-dir>/<same-basename>.docx and .pdf
Dependencies: python-docx, reportlab   ->  pip install python-docx reportlab
"""
import json
import os
import sys


def load(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


# ------------------------------- DOCX -------------------------------
def build_docx(d, out):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.45)
        s.bottom_margin = Inches(0.45)
        s.left_margin = Inches(0.55)
        s.right_margin = Inches(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)

    NAVY = RGBColor(0x1F, 0x3A, 0x5F)
    GREY = RGBColor(0x44, 0x44, 0x44)
    SOFT = RGBColor(0x55, 0x55, 0x55)

    def bottom_border(p):
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "999999")
        pbdr.append(bottom)
        pPr.append(pbdr)

    def section(title):
        p = doc.add_paragraph()
        r = p.add_run(title.upper())
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        bottom_border(p)

    # Name
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = name.add_run(d["name"])
    rn.bold = True
    rn.font.size = Pt(18)
    name.paragraph_format.space_after = Pt(0)

    # Contact
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = c.add_run("  |  ".join(d["contact"]))
    rc.font.size = Pt(9)
    rc.font.color.rgb = GREY
    c.paragraph_format.space_after = Pt(2)

    # Summary
    section("Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)
    sp.add_run(d["summary"]).font.size = Pt(9.5)

    # Skills
    section("Skills")
    for sk in d["skills"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        b = p.add_run(sk["category"] + ": ")
        b.bold = True
        b.font.size = Pt(9.5)
        p.add_run(sk["items"]).font.size = Pt(9.5)

    # Experience
    section("Experience")
    for e in d["experience"]:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(3)
        h.paragraph_format.space_after = Pt(0)
        hr = h.add_run(e["company"] + " — " + e["title"])
        hr.bold = True
        hr.font.size = Pt(10)
        meta = ", ".join([x for x in [e.get("dates"), e.get("location")] if x])
        if meta:
            m = h.add_run("   (" + meta + ")")
            m.italic = True
            m.font.size = Pt(9)
            m.font.color.rgb = SOFT
        for b in e["bullets"]:
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(1)
            bp.add_run(b).font.size = Pt(9.5)

    # Projects
    if d.get("projects"):
        section("Projects")
        for pr in d["projects"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            b = p.add_run(pr["name"] + ": ")
            b.bold = True
            b.font.size = Pt(9.5)
            p.add_run(pr["description"]).font.size = Pt(9.5)

    # Education
    if d.get("education"):
        section("Education")
        for ed in d["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            b = p.add_run(ed["institution"] + " — ")
            b.bold = True
            b.font.size = Pt(9.5)
            p.add_run(ed["detail"]).font.size = Pt(9.5)

    # Certifications
    if d.get("certifications"):
        section("Certifications")
        for ct in d["certifications"]:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            p.add_run(ct).font.size = Pt(9.5)

    doc.save(out)


# -------------------------------- PDF --------------------------------
def _esc(s):
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf(d, out):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, HRFlowable

    NAVY = colors.HexColor("#1F3A5F")
    GREY = colors.HexColor("#444444")

    name_st = ParagraphStyle("name", fontName="Helvetica-Bold", fontSize=18,
                             alignment=TA_CENTER, spaceAfter=2, leading=20)
    contact_st = ParagraphStyle("contact", fontName="Helvetica", fontSize=8.5,
                                alignment=TA_CENTER, textColor=GREY, spaceAfter=3, leading=10)
    head_st = ParagraphStyle("head", fontName="Helvetica-Bold", fontSize=10.5,
                             textColor=NAVY, spaceBefore=5, spaceAfter=1, leading=12)
    body_st = ParagraphStyle("body", fontName="Helvetica", fontSize=9,
                             spaceAfter=2, leading=11)
    exphead_st = ParagraphStyle("exphead", fontName="Helvetica-Bold", fontSize=10,
                                spaceBefore=3, spaceAfter=1, leading=11.5)
    bul_st = ParagraphStyle("bul", fontName="Helvetica", fontSize=9,
                            leftIndent=11, firstLineIndent=-7, spaceAfter=1, leading=11)

    flow = []

    def H(title):
        flow.append(Paragraph(_esc(title).upper(), head_st))
        flow.append(HRFlowable(width="100%", thickness=0.6,
                               color=colors.HexColor("#999999"),
                               spaceBefore=1, spaceAfter=2))

    flow.append(Paragraph(_esc(d["name"]), name_st))
    flow.append(Paragraph("  |  ".join(_esc(x) for x in d["contact"]), contact_st))

    H("Summary")
    flow.append(Paragraph(_esc(d["summary"]), body_st))

    H("Skills")
    for sk in d["skills"]:
        flow.append(Paragraph("<b>%s:</b> %s" % (_esc(sk["category"]), _esc(sk["items"])), body_st))

    H("Experience")
    for e in d["experience"]:
        meta = ", ".join([x for x in [e.get("dates"), e.get("location")] if x])
        head = "%s — %s" % (_esc(e["company"]), _esc(e["title"]))
        if meta:
            head += '  <font size=8.5 color="#555555"><i>(%s)</i></font>' % _esc(meta)
        flow.append(Paragraph(head, exphead_st))
        for b in e["bullets"]:
            flow.append(Paragraph("• " + _esc(b), bul_st))

    if d.get("projects"):
        H("Projects")
        for pr in d["projects"]:
            flow.append(Paragraph("<b>%s:</b> %s" % (_esc(pr["name"]), _esc(pr["description"])), body_st))

    if d.get("education"):
        H("Education")
        for ed in d["education"]:
            flow.append(Paragraph("<b>%s</b> — %s" % (_esc(ed["institution"]), _esc(ed["detail"])), body_st))

    if d.get("certifications"):
        H("Certifications")
        for ct in d["certifications"]:
            flow.append(Paragraph("• " + _esc(ct), bul_st))

    doc = SimpleDocTemplate(out, pagesize=LETTER, topMargin=0.45 * inch,
                            bottomMargin=0.4 * inch, leftMargin=0.55 * inch,
                            rightMargin=0.55 * inch, title=d["name"] + " - Resume")
    doc.build(flow)


def main():
    if len(sys.argv) < 2:
        print("usage: python build_resume.py <role.json>")
        sys.exit(1)
    jp = sys.argv[1]
    d = load(jp)
    base = os.path.splitext(jp)[0]
    build_docx(d, base + ".docx")
    build_pdf(d, base + ".pdf")
    print("Wrote:")
    print("  " + base + ".docx")
    print("  " + base + ".pdf")


if __name__ == "__main__":
    main()
